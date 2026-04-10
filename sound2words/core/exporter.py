from __future__ import annotations

from datetime import datetime
from pathlib import Path

from sound2words.core.models import SessionSnapshot


class TranscriptExporter:
    def __init__(self, project_root: Path) -> None:
        self._export_dir = project_root / "data" / "exports"
        self._export_dir.mkdir(parents=True, exist_ok=True)

    def export_txt(self, snapshot: SessionSnapshot) -> Path:
        path = self._build_path("txt")
        lines = [
            f"[{self._ms_to_hhmmss(s.start_ms)}-{self._ms_to_hhmmss(s.end_ms)}] {s.text}"
            for s in snapshot.segments
        ]
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def export_srt(self, snapshot: SessionSnapshot) -> Path:
        path = self._build_path("srt")
        blocks: list[str] = []
        for index, segment in enumerate(snapshot.segments, start=1):
            blocks.append(str(index))
            blocks.append(
                f"{self._ms_to_srt(segment.start_ms)} --> {self._ms_to_srt(segment.end_ms)}"
            )
            blocks.append(segment.text)
            blocks.append("")
        path.write_text("\n".join(blocks), encoding="utf-8")
        return path

    def export_docx(self, snapshot: SessionSnapshot) -> Path:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "未安装 python-docx，无法导出 Word。请在当前解释器执行: python -m pip install python-docx"
            ) from exc

        path = self._build_path("docx")
        doc = Document()
        doc.add_heading("语音转写结果", level=1)
        for segment in snapshot.segments:
            doc.add_paragraph(
                f"[{self._ms_to_hhmmss(segment.start_ms)}-{self._ms_to_hhmmss(segment.end_ms)}] {segment.text}"
            )
        doc.save(str(path))
        return path

    def _build_path(self, ext: str) -> Path:
        date_prefix = datetime.now().strftime("%Y-%m-%d")
        candidate = self._export_dir / f"{date_prefix}.{ext}"
        if not candidate.exists():
            return candidate

        index = 2
        while True:
            numbered = self._export_dir / f"{date_prefix}({index}).{ext}"
            if not numbered.exists():
                return numbered
            index += 1

    @staticmethod
    def _ms_to_hhmmss(value: int) -> str:
        total_seconds = value // 1000
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{hours:02}:{minutes:02}:{seconds:02}"

    @staticmethod
    def _ms_to_srt(value: int) -> str:
        total_ms = max(0, value)
        total_seconds, milliseconds = divmod(total_ms, 1000)
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{hours:02}:{minutes:02}:{seconds:02},{milliseconds:03}"
